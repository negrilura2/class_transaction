import pandas as pd
from docx import Document
from openai import OpenAI

# ========= 配置 =========

excel_file = "data.xlsx"
template_file = "template.docx"

client = OpenAI(
    api_key="你的API_KEY",
    base_url="https://api.deepseek.com"
)

# =======================


def generate_summary(name):

    prompt = f"""
请为一名软件工程专业大四学生写一段大学总结。

要求：
1 不超过150字
2 语气正式
3 内容包括学习、毕业设计、求职或实习
4 中文
5 每次生成内容不同

学生姓名：{name}
"""

    resp = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.9
    )

    text = resp.choices[0].message.content.strip()

    # 生成两段
    return text, text


def replace_in_paragraph(paragraph, data):
    for run in paragraph.runs:
        for key, value in data.items():
            if key in run.text:
                run.text = run.text.replace(key, value)


def replace_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:

            for p in cell.paragraphs:
                replace_in_paragraph(p, data)

            for t in cell.tables:
                replace_in_table(t, data)


df = pd.read_excel(excel_file)

for _, row in df.iterrows():

    sid = str(row.iloc[0])
    name = str(row.iloc[1])
    gpa = str(row.iloc[5])
    score = str(row.iloc[13])

    # 调用大模型
    text1, text2 = generate_summary(name)

    doc = Document(template_file)

    data = {
        "{{id}}": sid,
        "{{name}}": name,
        "{{gpa}}": gpa,
        "{{score}}": score
    }

    text_list = [text1, text2]
    text_index = 0

    # 替换段落
    for p in doc.paragraphs:

        if "{{text}}" in p.text:
            p.text = p.text.replace("{{text}}", text_list[text_index])
            text_index += 1
        else:
            replace_in_paragraph(p, data)

    # 替换表格
    for table in doc.tables:
        replace_in_table(table, data)

    output = f"{sid}.docx"
    doc.save(output)

    print("生成:", output)

print("全部完成")